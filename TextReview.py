import os
import re

import docx
from docx.shared import Pt
from docx.oxml.ns import qn


def extract_error_types_and_counts(doc_text):
    error_types = {}
    current_type = None

    for line in doc_text.split('\n'):
        line = line.strip()  # 去掉多余空白
        if re.match(r"错误类型\d+：", line):
            current_type = re.split(r"：", line, maxsplit=1)[-1].strip()
            if current_type not in error_types:
                error_types[current_type] = 0
        elif current_type and '原文：' in line:
            error_types[current_type] += 1

    return error_types


def update_summary(doc_path, out_path):
    os.makedirs('../files/out/Report', exist_ok=True)

    doc = docx.Document(doc_path)

    full_text = '\n'.join(paragraph.text for paragraph in doc.paragraphs)

    actual_error_counts = extract_error_types_and_counts(full_text)

    in_summary_section = False

    for paragraph in doc.paragraphs:
        if paragraph.text.startswith('总结'):
            in_summary_section = True
            continue

        if in_summary_section:
            if not paragraph.text.strip():
                break

            for error_type, count in actual_error_counts.items():
                pattern = rf"{error_type}：共出现\d+次"
                replacement = f"{error_type}：共出现{count}次"
                if re.search(pattern, paragraph.text):
                    paragraph.text = re.sub(pattern, replacement, paragraph.text)

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = '等线'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '等线')
            run.font.size = Pt(11)

    doc.save(os.path.join('../files/out/Report', os.path.basename(doc_path)))


if __name__ == '__main__':

    # 固定目录
    input_dir = "../files/LiveStreamerReport"
    output_dir = "../files/out/Report"

    print("请确认：需要处理的文件是否已放置到 '../files/LiveStreamerReport' 目录中？")
    user_input = input("输入 'y' 确认，或其他键退出程序: ").strip().lower()

    if user_input != 'y':
        print("操作已取消，请将文件放置到指定目录后再运行程序。")
        exit()

    # 检查输入目录是否存在
    if not os.path.exists(input_dir):
        print(f"输入目录不存在: {input_dir}")
        os.makedirs(input_dir)
        print("已自动创建输入目录，请将文件放置到该目录后重新运行程序。")
        exit()

    # 确保输出目录存在
    os.makedirs(output_dir, exist_ok=True)

    try:
        print("开始处理文档...")
        for filename in os.listdir(input_dir):
            if filename.endswith('.docx'):
                file_path = os.path.join(input_dir, filename)
                update_summary(file_path, output_dir)
        print(f"所有文件已处理完成！结果保存在目录: {output_dir}")
    except Exception as e:
        print(f"处理文档时发生错误: {e}")
    finally:
        input("按任意键退出程序...")
