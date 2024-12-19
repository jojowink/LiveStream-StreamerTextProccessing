import os
import re

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from openai import OpenAI

api_key = os.getenv("API_KEY")

input_folder = "./files/out"
output_folder = "./files/out/AI"

os.makedirs(output_folder, exist_ok=True)


def load_prompt_from_txt(file_path):
    """从 txt 文件中加载 Prompt"""
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"指定的 Prompt 文件不存在: {file_path}")
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read()


def process_text_with_ai(prompt, text, retries=3):
    """使用 OpenAI API 处理文本"""
    client = OpenAI(
        api_key=api_key,
        base_url="https://dashscope.aliyuncs.com/compatible-mode/v1",
    )

    for attempt in range(retries):
        try:
            response = client.chat.completions.create(
                model="qwen-plus",
                messages=[
                    {"role": "system", "content": prompt},
                    {"role": "user", "content": text}
                ]
            )
            return response.choices[0].message.content
        except Exception as e:
            print(f"处理请求出错：{e}，重试 {attempt + 1}/{retries}")
    print("重试次数用完，跳过当前请求。")
    return None


def process_doc_file(file_path, output_path, prompt):
    """处理单个 DOC 文件并保存结果"""
    try:
        doc = Document(file_path)
        full_text = "\n".join(paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip())
        ai_text = process_text_with_ai(prompt, full_text)
        if not ai_text:
            print(f"AI 未返回有效文本，跳过文件：{file_path}")
            return False

        new_doc = Document()
        new_paragraph = new_doc.add_paragraph(ai_text)
        for run in new_paragraph.runs:
            run.font.name = "等线"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '等线')
            run.font.size = Pt(11)
        new_doc.save(output_path)

        print(f"文件处理完成: {file_path}")
        return True
    except Exception as e:
        print(f"处理文件 {file_path} 发生错误: {e}")
        return False


def get_max_processed_number(output_folder):
    processed_files = [f for f in os.listdir(output_folder) if f.endswith((".doc", ".docx"))]
    max_number = 0

    for file in processed_files:
        numbers = re.findall(r'\d+', file)
        if numbers:
            max_number = max(max_number, int(numbers[-1]))  # 取数字的最大值

    return max_number


def log_failed_file(failed_file_path, failed_log_file="./files/log/failed_files.txt"):
    with open(failed_log_file, "a", encoding="utf-8") as log_file:
        log_file.write(f"{failed_file_path}\n")


def process_folder(input_folder, output_folder, prompt, failed_log_file="./files/log/failed_files.txt"):
    """处理文件夹中的所有 DOC 文件"""
    failed_files = []
    if os.path.exists(failed_log_file):
        with open(failed_log_file, "r", encoding="utf-8") as f:
            failed_files = [line.strip() for line in f.readlines() if line.strip()]
        print(f"检测到 {len(failed_files)} 个失败文件，优先重新处理...")

    if failed_files:
        for file_path in failed_files:
            if os.path.exists(file_path):
                output_path = os.path.join(output_folder, os.path.basename(file_path))
                success = process_doc_file(file_path, output_path, prompt)
                if success:
                    failed_files.remove(file_path)  # 成功后移除
                    print(f"重新处理成功: {file_path}")
                else:
                    print(f"重新处理失败: {file_path}")

        with open(failed_log_file, "w", encoding="utf-8") as f:
            f.writelines([f"{file}\n" for file in failed_files])
    max_processed_number = get_max_processed_number(output_folder)
    print(f"已处理的最大文件编号为: {max_processed_number}")

    all_files = [
        os.path.join(input_folder, f)
        for f in os.listdir(input_folder) if f.endswith((".doc", ".docx"))
    ]
    remaining_files = [
        f for f in all_files
        if
        any(part.isdigit() and int(part) > max_processed_number for part in
            re.split(r'(\d+)', os.path.basename(f)))
    ]

    remaining_files = sorted(
        remaining_files,
        key=lambda x: [int(part) if part.isdigit() else part for part in
                       re.split(r'(\d+)', os.path.basename(x))]
    )

    print(f"剩余待处理文件数量: {len(remaining_files)}")

    for file_path in remaining_files:
        output_path = os.path.join(output_folder, os.path.basename(file_path))
        success = process_doc_file(file_path, output_path, prompt)
        if not success:
            print(f"处理失败: {file_path}")
            log_failed_file(file_path, failed_log_file)  # 记录失败的文件


def main():
    input_folder_name = input("请输入需要处理的文件夹名称：")
    input_folder = os.path.join("./files/out/Text/", input_folder_name)
    output_folder = os.path.join("./files/out/AI/", input_folder_name)
    prompt_path = "./files/prompt/prompt.txt"

    if not os.path.exists(input_folder):
        print(f"错误：文件夹 {input_folder} 不存在！")
        return

    os.makedirs(output_folder, exist_ok=True)

    custom_prompt = load_prompt_from_txt(prompt_path)

    process_folder(input_folder, output_folder, custom_prompt)


if __name__ == '__main__':
    main()
    input("程序运行完成，按任意键退出程序...")
