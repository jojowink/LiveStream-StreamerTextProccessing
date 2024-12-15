import os
from docx import Document
from docx.shared import Pt  # 用于字体大小设置
from docx.oxml.ns import qn  # 用于中文字体支持
import re


def split_document_by_timestamp(input_file):
    doc = Document(input_file)

    base_name = os.path.splitext(os.path.basename(input_file))[0]

    output_dir = os.path.join('../files/out/Text', base_name)
    os.makedirs(output_dir, exist_ok=True)

    segments = []
    current_segment = ""

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if re.match(r'^\D+\d{2}:\d{2}:\d{2}', text):
            if current_segment:
                segments.append(current_segment)
            current_segment = text + "\n"
        else:
            current_segment += text + "\n"

    if current_segment:
        segments.append(current_segment)

    current_text = ""
    file_counter = 1

    for segment in segments:
        if len(current_text) + len(segment) > 8000:
            if current_text:
                output_file = os.path.join(output_dir, f"{base_name}_{file_counter}.docx")
                save_document(current_text, output_file)
                file_counter += 1
                current_text = segment
        else:
            current_text += segment

    if current_text:
        output_file = os.path.join(output_dir, f"{base_name}_{file_counter}.docx")
        save_document(current_text, output_file)


def save_document(text, output_file):
    doc = Document()
    for line in text.split("\n"):
        paragraph = doc.add_paragraph(line)

        for run in paragraph.runs:
            run.font.name = "Microsoft YaHei"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            run.font.size = Pt(11)
    doc.save(output_file)


def process_files():
    input_dir = '../files/LiveStreamerText'

    for filename in os.listdir(input_dir):
        if filename.endswith('.docx'):
            input_file = os.path.join(input_dir, filename)
            split_document_by_timestamp(input_file)


if __name__ == "__main__":
    try:
        # 提示用户确认文件是否已放置到指定目录
        print("请确认：需要处理的文件是否已放置到 '../files/LiveStreamerText' 目录中？")
        user_input = input("输入 'y' 确认，或其他键退出程序: ").strip().lower()

        if user_input == 'y':
            print("开始处理文件...")
            process_files()
            print("所有文件处理完成！输出结果已保存到 '../files/out/Text' 目录中。")
        else:
            print("操作已取消，请将文件放置到指定目录后再运行程序。")
    except Exception as e:
        print(f"程序运行出错: {e}")
    finally:
        input("按任意键退出程序...")
